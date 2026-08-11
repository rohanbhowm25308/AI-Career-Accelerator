"""
resume_parser.py
------------------
Extracts raw text from an uploaded resume (PDF / DOCX / TXT) and pulls out
useful structured info: contact details, detected skills and which of the
standard resume sections are present. Shared by every module in the
AI Career Accelerator (resume analyzer, interview engine, roadmap builder,
internship matcher all key off of `parse_resume`).
"""

import os
import re
import csv


try:
    import pdfplumber
except ImportError:  # pragma: no cover
    pdfplumber = None

try:
    import docx
except ImportError:  # pragma: no cover
    docx = None


# ---------------------------------------------------------------------------
# Skills database
# ---------------------------------------------------------------------------

def load_skills(csv_path="data/skills.csv"):
    """Load the master skills list from data/skills.csv -> list[str]."""
    skills = []
    if not os.path.exists(csv_path):
        return skills
    with open(csv_path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            skill = (row.get("skill") or "").strip().lower()
            if skill:
                skills.append(skill)
    return skills


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text(file_path):
    """Extract raw text from a resume file. Supports .pdf, .docx, .doc, .txt"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path):
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is not installed.")
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Default x_tolerance (3pt) merges adjacent words into one on
            # resume templates with tight character kerning (common with
            # Canva/design-tool exports) - e.g. "Machine Learning" comes out
            # as "MachineLearning", which then fails every downstream skill
            # match. A tighter tolerance reconstructs the real word breaks.
            page_text = page.extract_text(x_tolerance=1.5) or ""
            text_parts.append(page_text)
            # Resumes very often show "LinkedIn" / "GitHub" as a clickable
            # label with no visible URL in the text layer. Pull the actual
            # target URL out of the PDF's hyperlink annotations so contact
            # detection still finds it.
            try:
                for link in getattr(page, "hyperlinks", None) or []:
                    uri = link.get("uri")
                    if uri:
                        text_parts.append(uri)
            except Exception:
                pass
    return "\n".join(text_parts)


def _extract_docx(file_path):
    if docx is None:
        raise RuntimeError("python-docx is not installed.")
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]

    # Also pull text out of any tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)

    # Hyperlink targets (e.g. a "LinkedIn" label with no visible URL)
    try:
        for rel in document.part.rels.values():
            if "hyperlink" in rel.reltype and getattr(rel, "is_external", False):
                paragraphs.append(rel.target_ref)
    except Exception:
        pass

    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Contact info
# ---------------------------------------------------------------------------

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}")
LINKEDIN_URL_RE = re.compile(r"(linkedin\.com/[A-Za-z0-9\-_/%.]+)", re.IGNORECASE)
GITHUB_URL_RE = re.compile(r"(github\.com/[A-Za-z0-9\-_/%.]+)", re.IGNORECASE)
# Fallback for resumes that show a bare "LinkedIn" / "GitHub" label with the
# URL only embedded as a hyperlink target we couldn't extract (e.g. an
# unusual file format) -- still worth flagging as present.
LINKEDIN_LABEL_RE = re.compile(r"\blinkedin\b", re.IGNORECASE)
GITHUB_LABEL_RE = re.compile(r"\bgithub\b", re.IGNORECASE)


def extract_contact_info(text):
    email = EMAIL_RE.search(text)
    phone = PHONE_RE.search(text)
    linkedin = LINKEDIN_URL_RE.search(text) or LINKEDIN_LABEL_RE.search(text)
    github = GITHUB_URL_RE.search(text) or GITHUB_LABEL_RE.search(text)

    return {
        "email": email.group(0) if email else None,
        "phone": phone.group(0).strip() if phone else None,
        "linkedin": linkedin.group(0) if linkedin else None,
        "github": github.group(0) if github else None,
    }


# ---------------------------------------------------------------------------
# Skill extraction
# ---------------------------------------------------------------------------

def extract_skills(text, skills_list=None):
    """Return the subset of skills_list that appear in `text` (case-insensitive,
    whole-word / whole-phrase match)."""
    if skills_list is None:
        skills_list = load_skills()

    lowered = text.lower()
    # Some PDF exports (tight character kerning, certain design-tool
    # templates) drop the space between words in the extracted text layer,
    # e.g. "Machine Learning" -> "MachineLearning". Matching against a
    # whitespace-collapsed copy too means multi-word skills still get
    # detected even if that happens, independent of the extraction fix above.
    collapsed = re.sub(r"\s+", "", lowered)
    found = []

    for skill in skills_list:
        pattern = r"(?<![a-z0-9])" + re.escape(skill) + r"(?![a-z0-9])"
        if re.search(pattern, lowered):
            found.append(skill)
        elif " " in skill and skill.replace(" ", "") in collapsed:
            found.append(skill)

    return sorted(set(found))


# ---------------------------------------------------------------------------
# Section detection (used for the ATS heuristic score)
# ---------------------------------------------------------------------------

SECTION_KEYWORDS = {
    "experience": ["experience", "work history", "employment"],
    "education": ["education", "academic"],
    "skills": ["skills", "technical skills", "core competencies"],
    "projects": ["projects", "personal projects"],
    "summary": ["summary", "objective", "profile"],
    "certifications": ["certification", "certificate", "licenses"],
}


def detect_sections(text):
    lowered = text.lower()
    found = {}
    for section, keywords in SECTION_KEYWORDS.items():
        found[section] = any(kw in lowered for kw in keywords)
    return found


# ---------------------------------------------------------------------------
# Convenience: parse everything in one call
# ---------------------------------------------------------------------------

def parse_resume(file_path, skills_list=None):
    text = extract_text(file_path)
    word_count = len(text.split())

    return {
        "raw_text": text,
        "word_count": word_count,
        "contact_info": extract_contact_info(text),
        "skills": extract_skills(text, skills_list),
        "sections": detect_sections(text),
    }
